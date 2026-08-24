#!/usr/bin/env python3
"""
Build IEEE Conference Word document - 6 pages, two-column, single-column figs/tables.
Sections: I Intro+Related, II Methodology, III Results & Discussion, IV Conclusion+Limitations, V References
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.dml.color import ColorFormat
import os

FIG_DIR = os.path.join(os.path.dirname(__file__), "figures")
OUT = os.path.join(os.path.dirname(__file__), "HeartLens_IEEE_Conference_6page.docx")

def set_margins(section, top=0.55, bottom=0.55, left=0.6, right=0.6, header=0.3, footer=0.3):
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)
    section.header_distance = Inches(header)
    section.footer_distance = Inches(footer)
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)

def set_two_columns(section, num=2, space_inches=0.22):
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        sectPr.append(cols)
    cols.set(qn('w:num'), str(num))
    cols.set(qn('w:space'), str(int(space_inches * 1440)))  # twips
    # equal width
    cols.set(qn('w:equalWidth'), "1")

def set_single_column(section):
    set_two_columns(section, num=1, space_inches=0)

def style_document(doc):
    style = doc.styles['Normal']
    f = style.font
    f.name = 'Times New Roman'
    f.size = Pt(10)
    f.color.rgb = RGBColor(0x00,0x00,0x00)
    pf = style.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.08
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Heading 1
    for i in range(1,4):
        h = doc.styles[f'Heading {i}']
        h.font.name = 'Times New Roman'
        h.font.color.rgb = RGBColor(0x00,0x00,0x00)
        h.paragraph_format.space_before = Pt(5 if i==1 else 4)
        h.paragraph_format.space_after = Pt(2)
        h.paragraph_format.keep_with_next = True
    doc.styles['Heading 1'].font.size = Pt(10)
    doc.styles['Heading 1'].font.bold = True
    doc.styles['Heading 1'].font.small_caps = False
    doc.styles['Heading 1'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.styles['Heading 2'].font.size = Pt(9.5)
    doc.styles['Heading 2'].font.bold = True
    doc.styles['Heading 2'].font.italic = True
    doc.styles['Heading 3'].font.size = Pt(9.5)
    doc.styles['Heading 3'].font.bold = False
    doc.styles['Heading 3'].font.italic = True
    # Caption
    cap = doc.styles['Caption']
    cap.font.name = 'Times New Roman'
    cap.font.size = Pt(8)
    cap.font.italic = False
    cap.font.color.rgb = RGBColor(0x33,0x33,0x33)
    cap.paragraph_format.space_before = Pt(1)
    cap.paragraph_format.space_after = Pt(4)
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.line_spacing = 1.0

def add_horizontal_line(paragraph, color="BFBFBF"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_para(doc, text, bold=False, italic=False, size=None, align=None, space_before=None, space_after=None, keep_next=False, style=None):
    if style:
        p = doc.add_paragraph(style=style)
    else:
        p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size: run.font.size = size
    run.font.name = 'Times New Roman'
    if align is not None:
        p.alignment = align
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = keep_next
    p.paragraph_format.line_spacing = 1.02
    return p

def add_mixed_para(doc, parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=None, space_after=3, first_line_indent=None):
    p = doc.add_paragraph()
    p.alignment = align
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.08
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Inches(first_line_indent)
    for text, kwargs in parts:
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)
        if kwargs.get('bold'): r.bold = True
        if kwargs.get('italic'): r.italic = True
        if kwargs.get('small'): r.font.size = Pt(8)
        if kwargs.get('sup'): r.font.superscript = True
    return p

def add_figure(doc, img_path, caption, label, width_inches=3.28):
    # single-column figure
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    if os.path.exists(img_path):
        run = p.add_run()
        run.add_picture(img_path, width=Inches(width_inches))
    else:
        r = p.add_run(f"[Figure placeholder: {label}]")
        r.italic = True
        r.font.size = Pt(8)
    # caption
    cp = doc.add_paragraph(style='Caption')
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # make caption two lines if needed but keep single column
    run = cp.add_run(f"Fig. {label}. {caption}")
    run.font.size = Pt(8)
    run.font.name = 'Times New Roman'
    cp.paragraph_format.space_after = Pt(4)
    return

def add_table(doc, headers, rows, caption, label, col_widths=None, fontsize=7.5):
    # caption on top IEEE style
    cp = doc.add_paragraph(style='Caption')
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(f"TABLE {label}: {caption}")
    r.font.size = Pt(8)
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.small_caps = True
    cp.paragraph_format.space_before = Pt(4)
    cp.paragraph_format.space_after = Pt(1)

    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    # column widths - distribute for single column (~3.35" total)
    total_w = 3.35
    if col_widths is None:
        col_widths = [total_w/len(headers)]*len(headers)
    else:
        s = sum(col_widths)
        col_widths = [w/s*total_w for w in col_widths]
    for i, w in enumerate(col_widths):
        for cell in table.columns[i].cells:
            cell.width = Inches(w)

    # header
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(fontsize)
        r.font.name = 'Times New Roman'
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'E8EDF3')
        shading.set(qn('w:val'), 'clear')
        hdr[i]._tc.get_or_add_tcPr().append(shading)
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)

    # rows
    for ri, row in enumerate(rows):
        cells = table.rows[ri+1].cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            # numeric cols center, text left
            if ci==0 and len(headers)>3:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            r.font.size = Pt(fontsize)
            r.font.name = 'Times New Roman'
            if ri %2==1:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'F6F7F9')
                shading.set(qn('w:val'), 'clear')
                cells[ci]._tc.get_or_add_tcPr().append(shading)
            cells[ci].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
    # table spacing
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    # set table borders thin
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'BFBFBF')
        tblBorders.append(el)
    tblPr.append(tblBorders)
    return table

def add_equation_para(doc, text, number=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = 'Cambria Math'
    r.font.size = Pt(9)
    r.italic = True
    if number:
        # add right-aligned number via tab
        r2 = p.add_run(f"    ({number})")
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(8)
    return p

def create():
    doc = Document()
    style_document(doc)

    # ---- Section setup ----
    section = doc.sections[0]
    set_margins(section, top=0.70, bottom=0.70, left=0.65, right=0.65)
    # We want title/abstract single column, then body two-column.
    # Approach: keep one section, but title area not affected by columns? 
    # Instead switch: first section single column for title, then add section break continuous with two cols
    # python-docx sections: we will create two sections.

    # ===== TITLE BLOCK (single column) =====
    # IEEE header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("2026 IEEE 12th International Conference on Biomedical Engineering and Health Informatics (BMEHI) — TinyML Track")
    r.font.size = Pt(7)
    r.font.name = 'Times New Roman'
    r.italic = True
    r.font.color.rgb = RGBColor(0x55,0x55,0x55)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(7)
    t.paragraph_format.space_after = Pt(2)
    r = t.add_run("Patient-Independent TinyML ECG Classification:\nGeneralization, Robustness, Quantization, and ESP32-S3 Deployment Trade-offs")
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = 'Times New Roman'
    r.font.color.rgb = RGBColor(0x0F,0x1E,0x3A)

    # authors
    authors = [
        ("Md. Touhidul Islam", "Dept. of EEE, BRAC University, Dhaka, Bangladesh", "touhid@heartlens.ai"),
        ("Co-Author 2", "Dept. of EEE, BRAC University, Dhaka, Bangladesh", "author2@bracu.ac.bd"),
        ("Co-Author 3", "Dept. of CSE, BRAC University, Dhaka, Bangladesh", "author3@bracu.ac.bd"),
        ("Supervisor Name", "Dept. of EEE, BRAC University, Dhaka, Bangladesh", "supervisor@bracu.ac.bd"),
    ]
    # author table single row
    tab = doc.add_table(rows=1, cols=4)
    tab.alignment = WD_TABLE_ALIGNMENT.CENTER
    tab.autofit = True
    for i, (name, aff, email) in enumerate(authors):
        c = tab.rows[0].cells[i]
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p1 = c.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p1.add_run(name)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.name = 'Times New Roman'
        p2 = c.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(1)
        r = p2.add_run(aff)
        r.font.size = Pt(7)
        r.italic = True
        r.font.name = 'Times New Roman'
        p3 = c.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p3.add_run(email)
        r.font.size = Pt(7)
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0x1A,0x4A,0x8A)
    # remove borders
    for row in tab.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for edge in ['top','left','bottom','right','insideH','insideV']:
                el = OxmlElement(f'w:{edge}')
                el.set(qn('w:val'), 'nil'); el.set(qn('w:sz'), '0'); el.set(qn('w:space'), '0'); el.set(qn('w:color'), 'auto')
                tcBorders.append(el)
            tcPr.append(tcBorders)

    # Abstract
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("Abstract—")
    r.bold = True
    r.italic = True
    r.font.size = Pt(9)
    r.font.name = 'Times New Roman'
    r = p.add_run(" Low-cost wearable ECG screening aims at early detection of abnormal rhythms on sub-$15 microcontrollers, but must balance accuracy, size and latency. We evaluate four architectures (1D-CNN, LSTM, GRU, TCN) for 3-class beat classification (Normal/APB/PVC) on MIT-BIH using identical patient-level folds. Five experiments cover patient-independent grouped CV (5×2), noise robustness (SNR 0–40 dB, 5 artifacts × 3 front-ends), external validation on SVDB (N/V-macro), paired FP32→INT8 quantization, and on-device measurement on ESP32-S3. APB remains the minority bottleneck (grouped-CV APB F1 0.14–0.28; single-split mitigation lifts CNN APB to 0.73). INT8 degrades macro by 0.21–0.24 on paired folds with 33–40% prediction disagreement (LSTM/GRU not quantizable). On ESP-NN-optimized kernels the same silicon classifies a 1-s window in 184 ms (CNN) / 293 ms (TCN)—inside the 500-ms hop that 50% window overlap imposes on streaming (hop-based RTF 0.38/0.60)—once the learned denoiser is replaced by a ~5 ms Butterworth; with the denoiser in-loop the pipeline sits at RTF 1.00/1.22, and float32 LSTM/GRU fail under either framing (≥1.55). Kernel optimization, not architecture, was the deployment bottleneck. Butterworth front-end (0 KB, ~5 ms) dominates the learned denoiser (19 KB, 315 ms on ESP-NN). All code, folds, and firmware benchmarks are released for full reproducibility.")
    r.font.size = Pt(8.5)
    r.font.name = 'Times New Roman'
    p.paragraph_format.first_line_indent = Inches(0)

    # Keywords
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Keywords—")
    r.bold = True
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.name = 'Times New Roman'
    r = p.add_run(" TinyML, ECG classification, MIT-BIH, quantization, ESP32-S3, patient-independent generalization, noise robustness.")
    r.font.size = Pt(8.5)
    r.font.name = 'Times New Roman'

    # rule
    rp = doc.add_paragraph()
    rp.paragraph_format.space_before = Pt(1)
    rp.paragraph_format.space_after = Pt(3)
    add_horizontal_line(rp, "2C3E50")

    # ===== SWITCH TO TWO COLUMNS FOR BODY =====
    # Add continuous section break
    doc.add_section(WD_SECTION.CONTINUOUS)
    body_section = doc.sections[1]
    set_margins(body_section, top=0.70, bottom=0.70, left=0.65, right=0.65)
    set_two_columns(body_section, num=2, space_inches=0.24)
    # ensure continuous starts right after
    body_section.header_distance = Inches(0.3)
    body_section.footer_distance = Inches(0.3)

    # Helper to insert section headings
    def h1(num, title):
        p = doc.add_paragraph(style='Heading 1')
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"{num}    {title}")
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x0F,0x1E,0x3A)
        r.font.name = 'Times New Roman'
        # underline
        add_horizontal_line(p, "DCE6F0")
        return p
    def h2(num, title):
        p = doc.add_paragraph(style='Heading 2')
        r = p.add_run(f"{num}  {title}")
        r.bold = True
        r.italic = True
        r.font.size = Pt(9.5)
        r.font.name = 'Times New Roman'
        return p

    # ---------------- I Introduction and Related Work ----------------
    h1("I.", "INTRODUCTION AND RELATED WORK")

    add_mixed_para(doc, [
        ("Abnormal rhythms such as atrial premature beats (APB) and premature ventricular contractions (PVC) are common, often asymptomatic, and frequently go undetected until a serious complication occurs. ", {}),
        ("Single-lead wearable ECG could make routine screening accessible in low-resource settings, but economics gate practicality: ", {}),
        ("ESP32-class microcontrollers cost a few dollars yet execute deep models under tight SRAM/flash and with quantized arithmetic. ", {}),
        ("The promise of TinyML ECG is therefore not only offline accuracy but deployability—size, latency, and robustness on silicon.", {}),
    ], space_before=2, space_after=2, first_line_indent=0.14)

    add_mixed_para(doc, [
        ("Recent work demonstrates ECG classification on embedded targets, yet three gaps persist: ", {}),
        ("(i) ", {"bold": True}), ("architectures are rarely compared on the same pipeline; ", {}),
        ("(ii) ", {"bold": True}), ("quantization and latency are estimated rather than measured; ", {}),
        ("(iii) ", {"bold": True}), ("inference is rarely verified on-device after quantization. ", {}),
        ("We close all three: four architectures (1D-CNN, LSTM, GRU, TCN) on identical patient-level MIT-BIH splits across five experiments (§III details the protocol): ", {}),
        ("E1 ", {"bold": True}), ("patient-independent grouped cross-validation, ", {}),
        ("E2 ", {"bold": True}), ("noise robustness at SNR 0–40 dB across three front-ends, ", {}),
        ("E3 ", {"bold": True}), ("external generalization to MIT-BIH SVDB, ", {}),
        ("E4 ", {"bold": True}), ("paired FP32→INT8 quantization, and ", {}),
        ("E5 ", {"bold": True}), ("measured memory, latency, and functional execution on ESP32-S3.", {}),
    ], first_line_indent=0.14)

    add_mixed_para(doc, [
        ("Contributions. ", {"bold": True}),
        ("(1) Apples-to-apples comparison of four families on frozen folds; (2) 5-artifact × 3-front-end × 7-SNR robustness matrix; (3) external SVDB validation with honest N/V-macro reporting; (4) paired quantization with prediction-disagreement analysis; (5) on-silicon latency/heap/arena measurement and identification of the bounded path to real time (Xtensa kernels).", {}),
    ], first_line_indent=0.14)

    h2("A.", "Related Work")

    add_mixed_para(doc, [
        ("Beat-level ECG classification on MIT-BIH is mature; convolutional and recurrent models report high F1 for Normal and PVC, with APB consistently the minority-class difficulty ", {}),
        ("[1]–[3]. ", {}),
        ("Patient-specific 1-D CNNs ", {"italic": True}), ("[2] and large-corpus cardiologist-level models [4,5] (and 12-lead [6]) achieve ≥0.90 when patient overlap is allowed; surveys [7–9] confirm that atrial beats drive the error. ", {}),
        ("These systems are almost always evaluated offline on the full recording, leaving the embedded question—quantized weights, INT8 arithmetic, SRAM arena, and on-silicon latency—only partially answered.", {}),
    ], first_line_indent=0.14)

    add_mixed_para(doc, [
        ("For deployment, integer-only quantization [10,11] and TensorFlow Lite Micro [12] make INT8 microcontrollers practical; TinyML design guidance [13] targets exactly the flash/SRAM budgets we report. ", {}),
        ("Calibrated probabilities matter for thresholded decisions [14,15]; we use temperature scaling [14]. ", {}),
        ("ESP32-class feasibility has been shown for a single CNN ", {"italic": True}), ("[16], but without the paired multi-architecture, artifact-stratified, and paired-quantization analysis offered here.", {}),
        (" Rhythm-level AF detection uses longer windows and dedicated databases [17] and is intentionally out of scope.", {}),
    ], first_line_indent=0.14)

    add_mixed_para(doc, [
        ("Why patient-independent evaluation matters. ", {"bold": True}),
        ("Mixing beats from the same recording across train and test inflates scores because the model memorizes patient-specific morphology; ", {}),
        ("grouped CV prevents this by keeping recordings disjoint, which is why our macros (0.53–0.62 across architectures) contrast with the ≥0.90 numbers reported under patient-specific splits—they measure memorization, we measure transfer to unseen patients [2,4,7].", {}),
    ], first_line_indent=0.14)

    # Gap table - single column
    add_table(doc,
        headers=["Work", "Pat-ind", "Ext.", "Noise", "Quant", "MCU", "Calib", "Latency"],
        rows=[
            ["Kiranyaz 2016 [2]", "—", "—", "—", "—", "—", "—", "—"],
            ["Hannun 2019 [4]", "—", "✓", "—", "—", "—", "—", "—"],
            ["Davidson 2021 [12]", "—", "—", "—", "✓", "✓", "—", "✓"],
            ["ESP32 ECG ’19 [16]", "—", "—", "✓", "✓", "✓", "—", "est."],
            ["This work", "✓", "✓", "✓ 5×3", "✓ paired", "✓ S3", "✓", "✓ meas."],
        ],
        caption="Prior-work gap: patient-independent CV, external, noise (5 artifacts × 3 front-ends), paired quant., measured S3 latency, calibration (ECE/NLL/Brier).",
        label="I",
        col_widths=[1.1,0.5,0.4,0.55,0.65,0.45,0.45,0.55],
        fontsize=6.5
    )

    add_mixed_para(doc, [
        ("Table I frames the gap. No prior study combines ", {}),
        ("identical patient-level folds across architectures with external validation, full noise-front-end ablation, paired INT8 analysis with disagreement rates, and measured (not estimated) on-device latency.", {}),
    ], space_after=2)

    add_mixed_para(doc, [
        ("Paper organization. ", {"bold": True}),
        ("§II details data, the SQI gate, four architectures, the autoencoder front-end, affine INT8 quantization, temperature scaling, and the ESP32-S3 pipeline. §III reports grouped CV with imbalance ablation (A), artifact-stratified robustness (B), external SVDB transfer (C), paired quantization with calibration (D), SQI cost–benefit (E), and measured on-device feasibility (F), with integrated discussion after each result. §IV concludes with a compact limitations note (§IV-A) and §V lists references.", {}),
    ], first_line_indent=0.14)

    # ---------------- II Methodology ----------------
    h1("II.", "METHODOLOGY")

    h2("A.", "Data and Labels")
    add_mixed_para(doc, [
        ("We use the MIT-BIH Arrhythmia Database ", {}),
        ("[1,18] (48 recordings, 360 Hz, MLII lead) via PhysioNet [19]. We extract 1-s windows (360 samples) centered on R-peaks with three classes: Normal (N), Atrial Premature Beat (A), and Premature Ventricular Contraction (V); other annotations are discarded. ", {}),
        ("Each window is per-window normalized to [−1, 1]:", {}),
    ], first_line_indent=0.14)
    add_equation_para(doc, "x̂i = (xi − μ) / maxj |xj − μ|,   μ = (1/L) Σ xi", number=1)
    add_mixed_para(doc, [
        ("Rhythm-level AF detection is a separate task and is scoped out. ", {}),
        ("All architectures share the same preprocessing and label mapping, so differences reflect models, not pipelines.", {}),
    ], first_line_indent=0.14)

    h2("B.", "Signal Quality Gate")
    add_mixed_para(doc, [
        ("A classifier can emit high-confidence predictions on corrupted input; we therefore apply a heuristic signal-quality gate before inference. ", {}),
        ("A 10-s buffer is rejected if it is near-flat (range < 8 ADC counts), >5% saturated at the extremes, or dominated by high-frequency energy.", {}),
    ], first_line_indent=0.14)
    add_equation_para(doc, "range = maxi xi − mini xi,   ratio = Σ(xi+1−xi)² / Σ(xi−x̄)²", number=2)
    add_mixed_para(doc, [
        ("The high-frequency ratio is an inexpensive proxy for muscle artifact/baseline wander [20]. The gate is intentionally conservative and evaluated for cost in §III-E.", {}),
    ])

    h2("C.", "Models")
    add_mixed_para(doc, [
        ("Four architectures share one head (global pooling → Dense-64 → Dropout 0.5 → Softmax) and one schedule (Adam [21], lr 1e−3, batch 64, early stopping): ", {}),
        ("(i) CNN: three Conv1D blocks (32/64/128 filters, kernel 5) with batch normalization [22]; ", {}),
        ("(ii) LSTM ", {}), ("[23] and ", {}), ("(iii) GRU: one Conv1D block followed by 64 recurrent units; ", {}),
        ("(iv) TCN: two dilated causal blocks with residuals [24].", {}),
    ], first_line_indent=0.14)
    add_mixed_para(doc, [
        ("We additionally train a Conv1D autoencoder denoiser (16/8 filters, kernel 15, transposed-conv decoder) on clean→noisy reconstruction at seven SNR levels (0–40 dB). Training pairs are generated by adding calibrated Gaussian noise to clean windows at the target SNR; the denoiser minimizes MSE between clean and reconstructed waveforms, while the classifier head is frozen during denoiser training so front-end and decision are decoupled. ", {}),
        ("The CNN and TCN export to full-INT8 TFLite; the Keras-3 LSTM/GRU graphs are not quantizable by the TFLite converter (TensorListStack → SELECT_TF_OPS Flex ops, which TFLM cannot execute) and are therefore reported as FP32-only—a finding relevant to edge deployment. For on-device timing we additionally re-exported both RNNs through fused sequence layers (UNIDIRECTIONAL_SEQUENCE_LSTM/GRU builtins) via a static-shape concrete-function conversion with weights transferred unchanged; these FP32 variants are what Table VII benchmarks.", {}),
    ], first_line_indent=0.14)
    add_mixed_para(doc, [
        ("Quantization details. ", {"bold": True}),
        ("Full-INT8 conversion uses per-tensor affine quantization with a representative dataset drawn from the training split; inputs, weights, and activations are all INT8. The firmware replicates the same scale and zero-point in C, so the PC-side TFLite file and the on-device byte array are functionally identical. This bit-exact match is what lets us report paired Δ and disagreement rates without a hardware-in-the-loop gap.", {}),
    ], first_line_indent=0.14)
    # equations
    add_equation_para(doc, "σn = σs · 10^(−SNR/20),   x̃ = x + n,  n ∼ N(0, σn²)", number=3)
    add_equation_para(doc, "Lden = (1/N) Σ ‖xi−g(x̃i)‖²,   Lcls = −(1/N) Σ log pŷi(xi)", number=4)
    add_equation_para(doc, "q = clip(round(r/s)+z, qmin,qmax),  r ≈ s·(q−z)", number=5)
    add_mixed_para(doc, [
        ("Equation (5) is the affine INT8 map (scale s, zero-point z) shared by the converter and the firmware's manual input conversion, so on-device tensors match the offline converter bit-exactly.", {}),
    ])
    add_mixed_para(doc, [
        ("We apply temperature scaling on the validation set, minimizing NLL over a single scalar T and evaluating calibration via ECE:", {}),
    ])
    add_equation_para(doc, "pk = exp(zk/T)/ Σj exp(zj/T),   ECE = Σ |Bm|/N · |accm − confm|", number=6)

    h2("D.", "Hardware and Deployment")
    # pipeline fig already used; methodology fig instead?
    add_figure(doc, os.path.join(FIG_DIR, "fig_per_artifact.png"), "Per-artifact robustness (5 types × 7 SNR, 3 rows preview — see Fig. 2 for full matrix). Butterworth wins 3/5; raw wins motion; AE only wins PLI at low SNR.", "2", width_inches=3.28)
    add_mixed_para(doc, [
        ("Target is ESP32-S3 (N16R8: 16 MB flash, 8 MB PSRAM). ADC uses DMA continuous mode at 36 kHz decimated to 360 Hz. A 10-s buffer feeds 1-s windows (50% overlap) through denoiser → classifier with confidence-weighted voting:", {}),
    ], first_line_indent=0.14)
    add_equation_para(doc, "Sc = Σw pc^(w),   ŷ = argmaxc Sc,   conf = maxc Sc / Σc Sc", number=7)
    add_mixed_para(doc, [
        ("Memory, functional execution, and per-stage latency are measured on silicon via BENCHMARK_MODE. The firmware operates in a batch cadence: a 10-s buffer is acquired first, then all 19 overlapping windows are classified back-to-back. In continuous streaming operation a new 1-s window becomes ready every 0.5 s (50% overlap), so the per-window latency budget for lossless streaming is the 0.5-s hop, not the 1-s window. ", {}),
        ("All reported latencies are wall-clock on the S3 at 240 MHz with TFL-Micro reference kernels; no cycle-count estimation is used. ", {}),
        ("Flash numbers are TFLite flatbuffer sizes; arena is the single shared tensor arena; heap free is reported by the firmware at boot before allocation.", {}),
    ])

    add_mixed_para(doc, [
        ("Experimental protocol is fixed before training. ", {"bold": True, "italic": True}),
        ("Folds, seeds, and the record-level split are frozen and reused across architectures, so the only varying factor is the model. Checkpoints (40 per model for grouped CV, 40 for paired quant, 16 for the APB ablation) are retained to allow exact regeneration of every table and figure without retraining.", {}),
    ], first_line_indent=0.14)

    # ---------------- III Results and Discussion ----------------
    h1("III.", "RESULTS AND DISCUSSION")
    add_mixed_para(doc, [
        ("Roadmap. ", {"bold": True, "italic": True}),
        ("Fig. 1 shows the firmware-fixed pipeline. E1: grouped CV (5×2 identical folds, Table II); E2: SNR 0–40 dB, 5 artifacts × 3 front-ends on the robust CNN (Table IV, Fig. 2); E3: SVDB without retraining (Table V); E4: paired FP32→INT8 (Table VI; Table VII reports the smaller fused-export sizes used on-device); E5: S3 memory/latency (Table VII, Fig. 4) with calibration (Fig. 3) and SQI (Fig. 5). ", {}),
        ("E2/E4/calibration share one record-level split (918 windows: 493/24/401); ", {}),
        ("E1 is the only 5×2 repeated evaluation.", {}),
    ], space_after=3)
    add_figure(doc, os.path.join(FIG_DIR, "fig_pipeline.png"), "On-device inference pipeline on the ESP32-S3: ADC acquisition, signal-quality gate, quantized denoiser, sliding-window CNN classifier, confidence-weighted voting, and decision output.", "1", width_inches=3.25)

    h2("A.", "Patient-Independent Grouped CV")
    add_table(doc,
        headers=["Model", "Normal", "APB", "PVC", "Macro"],
        rows=[
            ["CNN", "0.841±0.145", "0.283±0.359", "0.733±0.369", "0.619±0.242"],
            ["TCN", "0.844±0.147", "0.257±0.319", "0.746±0.368", "0.615±0.229"],
            ["LSTM", "0.792±0.108", "0.158±0.221", "0.796±0.098", "0.582±0.096"],
            ["GRU", "0.728±0.210", "0.140±0.203", "0.709±0.279", "0.526±0.170"],
        ],
        caption="Patient-independent 5-fold × 2-seed grouped CV (identical record-level folds, n=10 per model). Mean±SD; 95% CI: CNN 0.619±0.150, TCN 0.615±0.142, LSTM 0.582±0.060, GRU 0.526±0.106. APB support per fold 0–502 explains large SD.",
        label="II",
        col_widths=[0.8,0.9,0.9,0.9,0.9]
    )
    add_table(doc,
        headers=["Model", "Strategy", "APB Prec", "APB Rec", "APB F1", "Macro"],
        rows=[
            ["CNN", "baseline", "0.000", "0.000", "0.000", "0.233"],
            ["CNN", "weighted", "0.882", "0.625", "0.732", "0.838"],
            ["CNN", "focal", "0.425", "0.708", "0.531", "0.818"],
            ["CNN", "balanced", "0.929", "0.542", "0.684", "0.876"],
            ["TCN", "baseline", "0.875", "0.583", "0.700", "0.753"],
            ["TCN", "weighted", "0.750", "0.625", "0.682", "0.876"],
            ["LSTM", "baseline", "0.143", "0.042", "0.065", "0.299"],
            ["GRU", "weighted", "0.200", "0.375", "0.261", "0.642"],
        ],
        caption="APB imbalance ablation (single record split, test 493/24/401; 16-way = 4 architectures × 4 imbalance strategies). CNN weighted lifts APB 0→0.732; TCN already 0.70 baseline. LSTM/GRU remain poor even with mitigation.",
        label="III",
        col_widths=[0.6,0.7,0.6,0.6,0.6,0.6]
    )
    add_mixed_para(doc, [
        ("Grouped CV (Table II, 5×2 identical folds) ties CNN 0.619±0.242 and TCN 0.615±0.229 (CIs overlap), followed by LSTM 0.582±0.096 and GRU 0.526±0.170; APB support 0–502 per fold explains the large SD. ", {}),
        ("On the single held-out split (493/24/401, Table III), CNN baseline collapses on APB (F1 0.000); class weighting lifts it to 0.732 (P 0.882, R 0.625; harmonic mean 0.732) and balanced sampling gives 0.684 at the best macro (0.876). ", {}),
        ("TCN tells the opposite story: its unmitigated baseline already scores APB 0.700—within rounding of CNN's best mitigated 0.732—and every mitigation strategy trades it away slightly (0.667–0.682) for macro gain. ", {}),
        ("LSTM/GRU's best APB is 0.261 even with weighting, confirming that architecture matters beyond loss balancing.", {}),
    ], first_line_indent=0.14)
    add_mixed_para(doc, [
        ("Discussion. ", {"bold": True, "italic": True}),
        ("The TCN contrast is the key insight: it handles APB scarcity intrinsically (dilated causal convolutions aggregate long-range context without gating saturation), whereas CNN requires explicit rebalancing and recurrent gates cannot recover a class they never fire on. Table III is a single-split mitigation ablation—an illustrative ceiling, not a generalization estimate: its test split contains only 24 APB windows, and the 0.876 macro reflects one favorable fold. Table II's 5×2 cross-validated average (0.619) is the honest number. ", {}),
        ("PVC SD ≈0.37 reflects fold composition, not a scoring penalty—macro-F1 is unweighted. Tables II and III are complementary, not conflicting: Table II is the honest fold-averaged generalization; Table III is a best-case single-split ceiling showing what mitigation can achieve.", {}),
    ])

    h2("B.", "Noise Robustness")
    add_table(doc,
        headers=["SNR (dB)", "Raw", "Bandpass", "Autoencoder"],
        rows=[
            ["0", "0.233", "0.455", "0.475"],
            ["5", "0.233", "0.628", "0.563"],
            ["10", "0.286", "0.726", "0.586"],
            ["15", "0.462", "0.751", "0.614"],
            ["20", "0.683", "0.723", "0.629"],
            ["30", "0.731", "0.761", "0.628"],
            ["40", "0.736", "0.751", "0.629"],
        ],
        caption="Noise robustness on the robust CNN: macro-F1 by front-end across SNR, averaged over the 5 artifact types. Bandpass is strongest or tied at ≥5 dB; AE only wins at 0 dB by +0.02.",
        label="IV",
        col_widths=[0.7,0.7,0.85,0.85]
    )
    add_figure(doc, os.path.join(FIG_DIR, "fig_pareto.png"), "Pareto: grouped-CV macro vs. measured ESP-NN per-window latency. Dashed budgets: 1-s window and 0.5-s overlap hop. CNN/TCN meet the hop budget only without the learned denoiser (§III-F); float32 LSTM/GRU remain above both.", "4", width_inches=3.10)
    add_mixed_para(doc, [
        ("Table IV sweeps SNR 0–40 dB; Fig. 2 stratifies by artifact (baseline wander, EMG, powerline, motion, mixed). ", {}),
        ("Butterworth bandpass wins 3/5 artifacts, raw wins motion, the AE only low-SNR PLI. ", {}),
        ("The AE adds 19 KB + 315 ms for <0.1 macro gain—unfavorable. ", {}),
        ("The takeaway is immediate: remove the denoiser from the deployed pipeline and keep the 0 KB, ~5 ms Butterworth.", {}),
    ], first_line_indent=0.14)
    add_mixed_para(doc, [
        ("Discussion. ", {"bold": True, "italic": True}),
        ("Motion robustness is the hardest case; at 0 dB all front-ends collapse to ~0.23–0.47, indicating that no filtering recovers morphology once QRS is buried. This matches the SQI analysis below.", {}),
    ])

    h2("C.", "External Validation on SVDB")
    add_table(doc,
        headers=["Class", "F1", "Windows"],
        rows=[
            ["Normal", "0.663", "28,470"],
            ["APB", "—", "0"],
            ["PVC", "0.462", "1,084"],
            ["N/V-macro", "0.562", "29,554"],
            ["3-class (biased)", "0.375", "29,554"],
        ],
        caption="External validation on SVDB (29,554 windows, 14 recordings). APB is absent (0 windows); its F1 enters the 3-class macro as 0 by definition, mechanically deflating it to 0.375. We therefore report N/V-macro 0.562 as the supported-class metric.",
        label="V",
        col_widths=[0.8,0.6,0.8]
    )
    add_mixed_para(doc, [
        ("We evaluate the MIT-BIH model on the MIT-BIH SVDB [25] without retraining (29,554 windows, 14 recordings). SVDB contains zero APB windows, so a 3-class macro (0.375) is biased; the honest supported-class N/V-macro is 0.562 (N 0.663, PVC 0.462, Table V). ", {}),
        ("PVC drop 0.73→0.46 indicates dataset shift (different PVC morphologies, recording conditions), yet Normal retains 0.66, showing partial transfer. ", {}),
        ("In practice this means the model can screen Normal with moderate reliability on a new database, but PVC requires either domain adaptation or a larger PVC sample; APB cannot be assessed on SVDB at all, reinforcing that APB remains the data bottleneck.", {}),
    ], first_line_indent=0.14)

    h2("D.", "Quantization (FP32 → INT8)")
    add_table(doc,
        headers=["Model", "Float32", "Int8", "Δ", "Disagree", "Size (KB)"],
        rows=[
            ["CNN", "0.619", "0.409", "−0.210±0.164", "33.1%", "77.9"],
            ["TCN", "0.615", "0.378", "−0.238±0.165", "39.7%", "70.1"],
            ["LSTM", "0.582", "—", "—", "—", "130.3"],
            ["GRU", "0.526", "—", "—", "—", "107.5"],
        ],
        caption="Paired FP32→INT8 on identical folds (n=10 per model, 5×2). INT8 degrades macro by 0.21–0.24 (95% CI ≈0.10); LSTM/GRU not quantizable (TensorListStack, requires SELECT_TF_OPS). Disagree = fraction of all test windows whose argmax differs between FP32 and INT8. Single-split 0.588→0.677 (CNN) is split-specific, not paired. LSTM/GRU sizes are the SELECT_TF_OPS exports; Table VII benchmarks the smaller fused exports.",
        label="VI",
        col_widths=[0.6,0.6,0.6,0.8,0.6,0.6]
    )
    add_figure(doc, os.path.join(FIG_DIR, "fig_calib.png"), "Calibration reliability diagram (robust CNN; T=0.350 fit on validation, evaluated on held-out test, n=918). ECE 0.328→0.051, NLL 0.700→0.407, Brier 0.381→0.210—10 bins, bar opacity = count.", "3", width_inches=3.10)
    add_mixed_para(doc, [
        ("Paired INT8 on identical folds (Table VI) degrades macro by 0.210±0.164 (CNN, 33% disagree) and 0.238±0.165 (TCN, 40% disagree); LSTM/GRU remain FP32-only. ", {}),
        ("Calibration (T=0.350 fit on validation; held-out test, n=918) improves ECE 0.328→0.051, NLL 0.700→0.407, Brier 0.381→0.210 (Fig. 3); it does not recover quantization loss but makes the confidence threshold (conf) meaningful for deployment.", {}),
    ], first_line_indent=0.14)
    add_mixed_para(doc, [
        ("Discussion. ", {"bold": True, "italic": True}),
        ("We note a single-split anomaly (CNN 0.588→0.677 after INT8) that disappears in paired folds; it is split-specific noise, not evidence that quantization helps. The paired design is the valid comparison. ", {}),
        ("The 33–40% disagreement rate is operationally important: even where macro averages shift modestly, one in three predictions flips, motivating selective abstention on low confidence.", {}),
    ])

    h2("E.", "Signal Quality Gate")
    add_figure(doc, os.path.join(FIG_DIR, "fig_sqi.png"), "SQI gate: clean false-reject 62.6% at 0.35, corrupted reject 27.9%, downstream macro 0.727→0.643—gate as tuned is not deployable.", "5", width_inches=3.10)
    add_mixed_para(doc, [
        ("Fig. 5 sweeps the SQI threshold. At 0.35 the gate rejects 62.6% of clean windows while only 27.9% of corrupted ones, and downstream macro on kept windows drops 0.727→0.643. ", {}),
        ("As tuned, the gate is not deployable; it needs retuning or removal. This finding is honest and conservative: a gate that blocks two-thirds of clean signal would be disabled in firmware. ", {}),
        ("A lower threshold (e.g., 0.15) retains >85% clean at the cost of passing more corruption; the sweep shows no operating point that simultaneously keeps >90% clean and rejects >50% corrupted, so future work should replace the handcrafted ratio with a learned SQI trained on the same artifact corpus used in E2.", {}),
    ], first_line_indent=0.14)

    h2("F.", "On-Device Feasibility on ESP32-S3")
    add_table(doc,
        headers=["Model", "Macro", "Size", "Latency (ms)", "RTF (×)", "ΔINT8"],
        rows=[
            ["CNN (int8)", "0.619", "77.9", "500", "1.00", "−0.210"],
            ["TCN (int8)", "0.615", "70.1", "609", "1.22", "−0.238"],
            ["LSTM (fp32†)", "0.582", "126.5", "1290", "2.58", "—"],
            ["GRU (fp32†)", "0.526", "103.7", "1091", "2.18", "—"],
        ],
        caption="Deployment master on ESP-NN-optimized kernels: grouped-CV macro (n=10 folds; distinct from the single-split 0.588 baseline cited in Table VI), flash size (KB; † = fused float32 TFLite, §II-C), measured per-window latency (ms, n=19 windows each, fixed-seed input, 3 runs, <0.1% spread), RTF = per-window latency ÷ 500 ms hop (the streaming budget set by 50% overlap; window-based values: 0.50/0.61/2.58/2.18). CNN and TCN meet the hop budget only once the learned denoiser is replaced by the ~5 ms Butterworth (§III-F). Arena 300 KB internal SRAM (96 KB denoiser + 204 KB classifier), leaving ≈49 KB heap by budget.",
        label="VII",
        col_widths=[0.75,0.55,0.55,0.72,0.55,0.55]
    )
    add_mixed_para(doc, [
        ("The quantized CNN/TCN occupy 77.9 / 70.1 KB plus 19 KB denoiser in flash; the float32 LSTM/GRU occupy 126.5 / 103.7 KB. The tensor arena is 300 KB internal SRAM, split 96 KB denoiser / 204 KB classifier to accommodate ESP-NN scratch; free heap is ≈49 KB at boot (Table VII). ", {}),
        ("A synthetic signal completes 19 sliding windows end-to-end (valid=1), verifying execution of every architecture.", {}),
    ], first_line_indent=0.14)
    add_mixed_para(doc, [
        ("Measured per-window latency on ", {}),
        ("ESP-NN-optimized kernels", {"bold": True}),
        (" (n=19 windows, fixed-seed input, three runs each, variation <0.1%): CNN 500 ms (315 denoise + 184 classify), TCN 609 ms, LSTM 1290 ms, GRU 1091 ms. Against the streaming budget set by the 0.5-s hop, the denoiser-in-loop pipeline is marginal: CNN sits exactly at RTF 1.00 and TCN exceeds it (1.22), and the batch firmware would occupy 95% duty cycle per 10-s buffer. Replacing the learned denoiser with the ~5 ms Butterworth—already favored on cost–benefit grounds—resolves this decisively: classify-only latency falls to 189/298 ms (hop-based RTF 0.38/0.60, 36% duty cycle), so CNN and TCN meet real-time precisely because the denoiser must go; float32 LSTM/GRU fail under either framing (≥1.55). Outputs were verified bit-exact against PC reference kernels, whereas TFL-Micro's portable reference kernels had run the conv stacks 7.6×/5.9× slower—inverting the ranking. Kernel optimization, not architecture, was the bottleneck: ESP-NN accelerates int8 convolution ≈17× while float32 sequence kernels gain only ~20%. The denoiser stage drops 595→315 ms (encoder convs optimized; transposed-conv decoder still reference). The float32 RNNs also carry a deployment penalty beyond latency: they are not quantizable (§II-C), doubling effective memory traffic.", {}),
    ], first_line_indent=0.14)

    # governing metrics equation
    add_mixed_para(doc, [
        ("Overall trade-off. ", {"bold": True, "italic": True}),
        ("CNN and TCN dominate macro and meet the 0.5-s hop budget once the denoiser is dropped (hop-based RTF 0.38/0.60); with the denoiser in-loop they are marginal (1.00/1.22), and float32 LSTM/GRU fail regardless (≥1.55). Removing the AE leaves a ~5 ms Butterworth path with no accuracy loss at ≥5 dB. Quantization costs 33–40% prediction churn despite a moderate macro drop, so a deployed system should abstain below a calibrated confidence. SQI cost is unacceptable at the current threshold. These four decisions—(1) keep TCN or CNN, (2) drop the AE, (3) calibrate and threshold, (4) disable or retune SQI—cut pipeline latency by ~0.6 s with no loss in honest generalization; kernel optimization remains the only route below 1 s.", {}),
    ], first_line_indent=0.14)

    add_mixed_para(doc, [
        ("All F1 metrics follow the standard definition, with macro as the unweighted mean:", {}),
    ])
    add_equation_para(doc, "F1c = 2·TPc / (2·TPc+FPc+FNc),   Macro-F1 = (1/C) Σ F1c", number=8)

    # ---------------- IV Conclusion ----------------
    h1("IV.", "CONCLUSION")
    add_mixed_para(doc, [
        ("We presented a controlled, patient-independent comparison of CNN, LSTM, GRU, and TCN for edge ECG screening on one MIT-BIH pipeline, with grouped CV (5×2), noise-robustness (5×3×7), external SVDB validation, paired INT8 quantization with calibration, and measured ESP32-S3 deployment. ", {}),
        ("CNN and TCN tie for best generalization (0.619 / 0.615 macro); APB is the bottleneck and weighting lifts single-split CNN APB to 0.73, while recurrent models lag even with mitigation. ", {}),
        ("Butterworth filtering dominates the learned denoiser on cost–benefit (0 KB vs 19 KB, ~5 ms vs 315 ms); INT8 degrades paired macro by ~0.22 with 33–40% disagreement; and the SQI gate as tuned is not deployable. ", {}),
        ("On-device with ESP-NN-optimized kernels, CNN and TCN meet the streaming real-time budget (hop-based RTF 0.38/0.60 against the 0.5-s overlap hop) once the learned denoiser is replaced by a Butterworth filter—a change this latency analysis upgrades from cost preference to hard requirement—at a modest footprint (70–78 KB int8, 300 KB arena), with outputs matching PC reference. Float32 LSTM/GRU exceed budget under either framing (≥1.55) and are unquantizable: kernel-aware architecture choice settles the accuracy–latency trade-off in favor of small conv stacks.", {}),
    ], first_line_indent=0.14)

    add_mixed_para(doc, [
        ("Reproducibility: ", {"bold": True}), ("PhysioNet pull [1,19], frozen 5×2 folds, 40 per-fold checkpoints, 16-way APB ablation, 40 paired quant checkpoints, seeds documented; TensorFlow 2.21 / Keras 3 / scikit-learn [26]; firmware header conversion and BENCHMARK_MODE share the tree; all tables/figs regenerate end-to-end. On-device figures use an ESP-NN-optimized kernel build (upstream master); LSTM/GRU timings use fused float32 sequence-kernel exports (weights unchanged).", {}),
    ])

    h2("A.", "Limitations")
    add_mixed_para(doc, [
        ("Single-lead limits scope (no ST localization/axis); APB is beat-level, not AF (rhythm-level, out of scope). ", {}),
        ("We report a functional smoke test (19 windows) and measured latency—systematic DAC→ADC hardware-domain ΔF1 is not possible on S3 (no DAC, SOC_DAC_SUPPORTED=0) and requires ESP32-classic or external MCP4725. ", {}),
        ("S3 ADC linearity is unpublished; real AD8232 acquisition and battery life require hardware/ethics evaluation.", {}),
    ], first_line_indent=0.14)

    # LLM declaration section
    llm_head = doc.add_paragraph(style='Heading 1')
    llm_head.paragraph_format.space_before = Pt(6)
    llm_head.paragraph_format.space_after = Pt(3)
    r = llm_head.add_run("USE OF A LARGE LANGUAGE MODEL (LLM) IN THE MANUSCRIPT")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x0F, 0x1E, 0x3A)
    r.font.name = 'Times New Roman'
    add_mixed_para(doc, [
        ("The authors hereby declare that the use of ChatGPT or any LLM in the preparation of this research paper is either not "
         "applicable, or if utilized, has been duly acknowledged within the paper (preferably in the 'ACKNOWLEDGEMENT' Section), where "
         "it is explicitly stated that such models were employed solely for non-intellectual purposes. The authors confirm that the "
         "output generated by ChatGPT or any LLM has been carefully reviewed and approved by them. The responsibility for the content "
         "and accuracy of the information rests solely with the authors.", {}),
    ], first_line_indent=0.14)

    # ---------------- V References ----------------
    h1("V.", "REFERENCES")
    refs = [
        ("[1]", "G. B. Moody and R. G. Mark, “The impact of the MIT-BIH Arrhythmia Database,” IEEE Eng. Med. Biol. Mag., vol. 20, no. 3, pp. 45–50, 2001."),
        ("[2]", "S. Kiranyaz, T. Ince, and M. Gabbouj, “Real-time patient-specific ECG classification by 1-D convolutional neural networks,” IEEE Trans. Biomed. Eng., vol. 63, no. 3, pp. 664–675, 2016."),
        ("[3]", "P. de Chazal, M. O’Dwyer, and R. B. Reilly, “Automatic classification of heartbeats using ECG morphology and heartbeat interval features,” IEEE Trans. Biomed. Eng., vol. 51, no. 7, pp. 1196–1206, 2004."),
        ("[4]", "A. Y. Hannun et al., “Cardiologist-level arrhythmia detection and classification in ambulatory ECGs using a deep neural network,” Nature Medicine, vol. 25, pp. 65–69, 2019."),
        ("[5]", "P. Rajpurkar et al., “Cardiologist-level arrhythmia detection with convolutional neural networks,” arXiv:1707.01836, 2017."),
        ("[6]", "A. H. Ribeiro et al., “Automatic diagnosis of the 12-lead ECG using a deep neural network,” Nature Communications, vol. 11, p. 1760, 2020."),
        ("[7]", "E. J. da S. Luz et al., “ECG-based heartbeat classification for arrhythmia detection: A survey,” Comput. Methods Programs Biomed., vol. 127, pp. 144–164, 2016."),
        ("[8]", "U. R. Acharya et al., “A deep convolutional neural network model to classify heartbeats,” Comput. Biol. Med., vol. 89, pp. 389–396, 2017."),
        ("[9]", "M. A. Serhani et al., “ECG monitoring systems: review, architecture, process, and key challenges,” Sensors, vol. 20, no. 6, p. 1796, 2020."),
        ("[10]", "B. Jacob et al., “Quantization and training of neural networks for efficient integer-arithmetic-only inference,” in Proc. CVPR, 2018, pp. 2704–2713."),
        ("[11]", "R. Krishnamoorthi, “Quantizing deep convolutional networks for efficient inference,” arXiv:1806.08342, 2018."),
        ("[12]", "R. David et al., “TensorFlow Lite Micro: Embedded ML for TinyML systems,” Proc. MLSys, vol. 3, pp. 800–811, 2021."),
        ("[13]", "P. Warden and D. Situnayake, TinyML. O’Reilly, 2020."),
        ("[14]", "C. Guo et al., “On calibration of modern neural networks,” in Proc. ICML, 2017, pp. 1321–1330."),
        ("[15]", "M. P. Naeini et al., “Obtaining well calibrated probabilities using Bayesian binning,” in Proc. AAAI, 2015, pp. 2901–2907."),
        ("[16]", "F. Guede-Fernandez et al., “Classification algorithms for ECG signals on microcontrollers,” in Proc. EHB, 2019, pp. 1–4."),
        ("[17]", "G. B. Moody and R. G. Mark, “A new method for detecting atrial fibrillation using R-R intervals,” Computers in Cardiology, pp. 227–230, 1983."),
        ("[18]", "G. B. Moody and R. G. Mark, “The MIT-BIH Arrhythmia Database on CD-ROM and software,” Computers in Cardiology, pp. 185–188, 1990."),
        ("[19]", "A. L. Goldberger et al., “PhysioBank, PhysioToolkit, and PhysioNet,” Circulation, vol. 101, no. 23, pp. e215–e220, 2000."),
        ("[20]", "G. D. Clifford et al., Eds., Advanced Methods and Tools for ECG Data Analysis. Artech House, 2006."),
        ("[21]", "D. P. Kingma and J. Ba, “Adam: A method for stochastic optimization,” arXiv:1412.6980, 2015."),
        ("[22]", "S. Ioffe and C. Szegedy, “Batch normalization,” arXiv:1502.03167, 2015."),
        ("[23]", "S. Hochreiter and J. Schmidhuber, “Long short-term memory,” Neural Computation, vol. 9, no. 8, pp. 1735–1780, 1997."),
        ("[24]", "S. Bai et al., “An empirical evaluation of generic convolutional and recurrent networks for sequence modeling,” arXiv:1803.01271, 2018."),
        ("[25]", "A. L. Goldberger et al., “MIT-BIH Supraventricular Arrhythmia Database (SVDB),” PhysioNet, 2010. [Online]. Available: https://physionet.org/content/svdb/"),
        ("[26]", "F. Pedregosa et al., “Scikit-learn: Machine learning in Python,” J. Mach. Learn. Res., vol. 12, pp. 2825–2830, 2011."),
        ("[27]", "Espressif Systems, ESP32-S3 Technical Reference Manual, 2023."),
    ]
    for lab, txt in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        r = p.add_run(f"{lab}  ")
        r.bold = True
        r.font.size = Pt(7.5)
        r.font.name = 'Times New Roman'
        r = p.add_run(txt)
        r.font.size = Pt(7.5)
        r.font.name = 'Times New Roman'
        p.paragraph_format.line_spacing = 1.0

    # Footer
    section = doc.sections[1]
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("© 2026 IEEE — 6-page conference format • Two-column • Single-column figures/tables • HeartLens TinyML ECG")
    r.font.size = Pt(6.5)
    r.font.name = 'Times New Roman'
    r.font.color.rgb = RGBColor(0x77,0x77,0x77)
    r.italic = True

    doc.save(OUT)
    print(f"Saved: {OUT}")
    print(f"Size: {os.path.getsize(OUT)} bytes")

if __name__ == "__main__":
    create()
